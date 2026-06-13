import os
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from src.logger import logger
from src.config import CHUNK_SIZE, CHUNK_OVERLAP

try:
    import pypdf
except ImportError:
    pypdf = None
    logger.warning("pypdf is not installed. PDF parsing will fall back to basic text matching if applicable.")

try:
    import docx
except ImportError:
    docx = None
    logger.warning("python-docx is not installed. Word document parsing will not be supported.")


class DocumentParser:
    """Parses various document types (PDF, DOCX, TXT) into raw text."""

    @staticmethod
    def parse_pdf(file_path: Path) -> str:
        """Extracts text from a PDF file using pypdf."""
        if not pypdf:
            raise ImportError("pypdf library is required to parse PDF files. Run 'pip install pypdf'.")
        
        text_content = []
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {e}")
            raise

    @staticmethod
    def parse_docx(file_path: Path) -> str:
        """Extracts text from a DOCX file using python-docx."""
        if not docx:
            raise ImportError("python-docx library is required to parse DOCX files. Run 'pip install python-docx'.")
        
        try:
            doc = docx.Document(file_path)
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise

    @staticmethod
    def parse_txt(file_path: Path) -> str:
        """Extracts text from a plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {e}")
            raise

    def parse(self, file_path: str | Path) -> str:
        """Determines file type and parses it accordingly."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        logger.info(f"Parsing file: {path.name} with extension: {suffix}")

        if suffix == ".pdf":
            return self.parse_pdf(path)
        elif suffix == ".docx":
            return self.parse_docx(path)
        elif suffix in [".txt", ".md"]:
            return self.parse_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")


def clean_text(text: str) -> str:
    """Cleans extracted text by normalizing whitespace, removing noise, etc."""
    # Replace multiple newlines/whitespaces with a single space
    text = re.sub(r"\s+", " ", text)
    # Remove control characters
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]", "", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Splits a document's clean text into overlapping chunks."""
    words = text.split(" ")
    chunks = []
    
    # Simple word-based chunker
    i = 0
    while i < len(words):
        chunk_words = words[i : i + chunk_size]
        chunks.append(" ".join(chunk_words))
        # Move forward by chunk_size - chunk_overlap
        i += max(1, chunk_size - chunk_overlap)
        
    return chunks


def extract_metadata(text: str) -> Dict[str, Any]:
    """Extracts basic info like email, phone, and name candidate from text."""
    metadata = {}
    
    # Extract Email
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    metadata["email"] = email_match.group(0) if email_match else None
    
    # Extract Phone (simple regex for various formats)
    phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    metadata["phone"] = phone_match.group(0) if phone_match else None
    
    return metadata
